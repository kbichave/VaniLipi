"""
Transcript export service.
Converts transcript segments to SRT, VTT, TXT, DOCX, PDF, and JSON.

All functions accept a list of segment dicts:
    {"start": float, "end": float, "marathi": str, "english": str}
"""
from __future__ import annotations

import json
import io
from typing import Any


# ─── Helpers ──────────────────────────────────────────────────────────────────


def _srt_ts(secs: float) -> str:
    """Convert seconds to SRT timestamp: HH:MM:SS,mmm"""
    h = int(secs // 3600)
    m = int((secs % 3600) // 60)
    s = int(secs % 60)
    ms = int((secs % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _vtt_ts(secs: float) -> str:
    """Convert seconds to VTT timestamp: HH:MM:SS.mmm"""
    return _srt_ts(secs).replace(",", ".")


# ─── Exporters ────────────────────────────────────────────────────────────────


def to_srt(segments: list[dict[str, Any]]) -> bytes:
    """Export to SRT subtitle format."""
    lines = []
    for i, seg in enumerate(segments, 1):
        lines.append(str(i))
        lines.append(f"{_srt_ts(seg['start'])} --> {_srt_ts(seg['end'])}")
        if seg.get("marathi"):
            lines.append(seg["marathi"])
        if seg.get("english"):
            lines.append(seg["english"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def to_vtt(segments: list[dict[str, Any]]) -> bytes:
    """Export to WebVTT format."""
    lines = ["WEBVTT", ""]
    for i, seg in enumerate(segments, 1):
        lines.append(f"cue-{i}")
        lines.append(f"{_vtt_ts(seg['start'])} --> {_vtt_ts(seg['end'])}")
        if seg.get("marathi"):
            lines.append(seg["marathi"])
        if seg.get("english"):
            lines.append(seg["english"])
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def to_txt(segments: list[dict[str, Any]], mode: str = "both") -> bytes:
    """
    Export to plain text.
    mode: "both" | "marathi" | "english"
    """
    lines = []
    for seg in segments:
        ts = f"[{_fmt_ts(seg['start'])}]"
        if mode in ("both", "marathi") and seg.get("marathi"):
            lines.append(f"{ts} {seg['marathi']}")
        if mode in ("both", "english") and seg.get("english"):
            lines.append(f"     {seg['english']}")
        if mode == "both":
            lines.append("")
    return "\n".join(lines).encode("utf-8")


def _fmt_ts(secs: float) -> str:
    """MM:SS for plain text."""
    m = int(secs // 60)
    s = int(secs % 60)
    return f"{m:02d}:{s:02d}"


def to_json(segments: list[dict[str, Any]]) -> bytes:
    """Export to structured JSON."""
    payload = {
        "transcript": [
            {
                "index": i,
                "start": seg["start"],
                "end": seg["end"],
                "marathi": seg.get("marathi", ""),
                "english": seg.get("english", ""),
            }
            for i, seg in enumerate(segments)
        ]
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def to_docx(segments: list[dict[str, Any]]) -> bytes:
    """Export to DOCX. Requires python-docx."""
    try:
        from docx import Document  # type: ignore[import]
        from docx.shared import Pt, RGBColor  # type: ignore[import]
    except ImportError as exc:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx") from exc

    doc = Document()
    doc.add_heading("VaniLipi Transcript", level=1)

    for seg in segments:
        ts = _fmt_ts(seg["start"])
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{ts}]  ")
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        if seg.get("marathi"):
            mar_run = p.add_run(seg["marathi"])
            mar_run.bold = True
            mar_run.font.size = Pt(12)

        if seg.get("english"):
            doc.add_paragraph(seg["english"]).runs[0].font.size = Pt(10) if doc.paragraphs else None  # type: ignore[assignment]
            # Re-add properly
            eng_p = doc.paragraphs[-1]
            eng_p.paragraph_format.left_indent = Pt(36)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(segments: list[dict[str, Any]]) -> bytes:
    """Export to PDF. Requires fpdf2."""
    try:
        from fpdf import FPDF  # type: ignore[import]
    except ImportError as exc:
        raise RuntimeError("fpdf2 not installed. Run: pip install fpdf2") from exc

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "VaniLipi Transcript", ln=True)
    pdf.ln(4)

    for seg in segments:
        ts = _fmt_ts(seg["start"])
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(148, 163, 184)
        pdf.cell(0, 5, f"[{ts}]", ln=True)

        if seg.get("marathi"):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(26, 26, 46)
            # PDF doesn't render Devanagari without a font file; fall back to latin placeholder
            safe_text = seg["marathi"].encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe_text)

        if seg.get("english"):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(74, 85, 104)
            pdf.multi_cell(0, 6, seg["english"])

        pdf.ln(2)

    return pdf.output()


# ─── Dispatch ─────────────────────────────────────────────────────────────────

EXPORTERS: dict[str, tuple[Any, str]] = {
    "srt":  (to_srt,  "text/plain"),
    "vtt":  (to_vtt,  "text/vtt"),
    "txt":  (to_txt,  "text/plain"),
    "json": (to_json, "application/json"),
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf":  (to_pdf,  "application/pdf"),
}


def export(
    segments: list[dict[str, Any]],
    fmt: str,
) -> tuple[bytes, str]:
    """
    Export segments to the requested format.
    Returns (content_bytes, mime_type).
    Raises ValueError for unknown formats, RuntimeError if a dep is missing.
    """
    if fmt not in EXPORTERS:
        raise ValueError(f"Unknown export format: {fmt!r}. Supported: {list(EXPORTERS)}")
    fn, mime = EXPORTERS[fmt]
    return fn(segments), mime
